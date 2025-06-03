"""
Document parsing utilities for FDA Regulatory Assistant.

This module provides robust document parsing functionality for various file formats
including PDFs, Word documents, and text files. It uses PyMuPDF for PDF parsing
and python-docx for Word document parsing.
"""

import io
import logging
from typing import Dict, Any, Optional
import re

# Try to import PyMuPDF with fallback
try:
    import pymupdf as fitz  # Use pymupdf import to avoid namespace conflicts

    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import fitz  # Fallback to fitz import

        PYMUPDF_AVAILABLE = True
    except ImportError:
        PYMUPDF_AVAILABLE = False
        print("Warning: PyMuPDF not available. PDF parsing will be limited.")

# Try to import python-docx with fallback
try:
    from docx import Document

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word document parsing will be limited.")

# Set up logging
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Document parser class that handles multiple file formats.

    Supports:
    - PDF files (using PyMuPDF/fitz)
    - Word documents (.docx, .doc)
    - Text files (.txt)
    - Other text-based formats
    """

    def __init__(self):
        """Initialize the document parser."""
        self.supported_formats = {
            ".txt": self._parse_text,
        }

        # Add PDF support if PyMuPDF is available
        if PYMUPDF_AVAILABLE:
            self.supported_formats[".pdf"] = self._parse_pdf

        # Add Word document support if python-docx is available
        if PYTHON_DOCX_AVAILABLE:
            self.supported_formats[".docx"] = self._parse_docx
            self.supported_formats[".doc"] = (
                self._parse_docx
            )  # python-docx can handle .doc files too

    def extract_text_from_file(
        self, file_content: bytes, filename: str
    ) -> Dict[str, Any]:
        """
        Extract text content from uploaded files with metadata.

        Args:
            file_content: Raw file content as bytes
            filename: Name of the file for format detection

        Returns:
            Dictionary containing:
            - text: Extracted text content
            - metadata: File metadata (pages, word_count, etc.)
            - success: Whether parsing was successful
            - error: Error message if parsing failed
        """
        try:
            # Get file extension
            file_ext = self._get_file_extension(filename)

            # Check if format is supported
            if file_ext not in self.supported_formats:
                # Try to decode as text for unknown formats
                return self._parse_text(file_content, filename)

            # Parse using appropriate parser
            parser_func = self.supported_formats[file_ext]
            return parser_func(file_content, filename)

        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            return {
                "text": f"[Error extracting content from {filename}: {str(e)}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "pages": 0,
                    "word_count": 0,
                    "parsing_error": str(e),
                },
                "success": False,
                "error": str(e),
            }

    def _get_file_extension(self, filename: str) -> str:
        """Get normalized file extension with dot."""
        if "." in filename:
            return "." + filename.lower().split(".")[-1]
        return ""

    def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF files using PyMuPDF.

        Args:
            file_content: PDF file content as bytes
            filename: Name of the PDF file

        Returns:
            Dictionary with extracted text and metadata
        """
        if not PYMUPDF_AVAILABLE:
            return {
                "text": f"[PDF parsing not available - PyMuPDF not installed. File: {filename}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "pages": 0,
                    "word_count": 0,
                    "file_type": "PDF",
                    "parsing_error": "PyMuPDF not available",
                },
                "success": False,
                "error": "PyMuPDF not available",
            }

        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=file_content, filetype="pdf")

            extracted_text = []
            page_count = len(pdf_document)

            # Extract text from each page
            for page_num in range(page_count):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()

                if page_text.strip():  # Only add non-empty pages
                    # Add page header for reference
                    extracted_text.append(f"\n--- Page {page_num + 1} ---\n")
                    extracted_text.append(page_text)

            # Close the document
            pdf_document.close()

            # Combine all text
            full_text = "".join(extracted_text)

            # Clean up the text
            full_text = self._clean_extracted_text(full_text)

            # Calculate metadata
            word_count = len(full_text.split()) if full_text else 0

            return {
                "text": full_text,
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "pages": page_count,
                    "word_count": word_count,
                    "file_type": "PDF",
                    "parser": "PyMuPDF",
                },
                "success": True,
                "error": None,
            }

        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {str(e)}")
            return {
                "text": f"[Error parsing PDF {filename}: {str(e)}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "pages": 0,
                    "word_count": 0,
                    "file_type": "PDF",
                    "parsing_error": str(e),
                },
                "success": False,
                "error": str(e),
            }

    def _parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Word documents using python-docx.

        Args:
            file_content: Word document content as bytes
            filename: Name of the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        if not PYTHON_DOCX_AVAILABLE:
            return {
                "text": f"[Word document parsing not available - python-docx not installed. File: {filename}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "paragraphs": 0,
                    "word_count": 0,
                    "file_type": "Word Document",
                    "parsing_error": "python-docx not available",
                },
                "success": False,
                "error": "python-docx not available",
            }

        try:
            # Open document from bytes
            doc = Document(io.BytesIO(file_content))

            extracted_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text.append(" | ".join(row_text))

            # Combine all text
            full_text = "\n".join(extracted_text)

            # Clean up the text
            full_text = self._clean_extracted_text(full_text)

            # Calculate metadata
            word_count = len(full_text.split()) if full_text else 0
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])

            return {
                "text": full_text,
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "paragraphs": paragraph_count,
                    "word_count": word_count,
                    "file_type": "Word Document",
                    "parser": "python-docx",
                },
                "success": True,
                "error": None,
            }

        except Exception as e:
            logger.error(f"Error parsing Word document {filename}: {str(e)}")
            return {
                "text": f"[Error parsing Word document {filename}: {str(e)}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "paragraphs": 0,
                    "word_count": 0,
                    "file_type": "Word Document",
                    "parsing_error": str(e),
                },
                "success": False,
                "error": str(e),
            }

    def _parse_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse text files and other text-based formats.

        Args:
            file_content: Text file content as bytes
            filename: Name of the text file

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Try different encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                # Fallback: decode with errors='ignore'
                text = file_content.decode("utf-8", errors="ignore")
                used_encoding = "utf-8 (with errors ignored)"

            # Clean up the text
            text = self._clean_extracted_text(text)

            # Calculate metadata
            word_count = len(text.split()) if text else 0
            line_count = len(text.splitlines()) if text else 0

            return {
                "text": text,
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "lines": line_count,
                    "word_count": word_count,
                    "file_type": "Text",
                    "encoding": used_encoding,
                },
                "success": True,
                "error": None,
            }

        except Exception as e:
            logger.error(f"Error parsing text file {filename}: {str(e)}")
            return {
                "text": f"[Error parsing text file {filename}: {str(e)}]",
                "metadata": {
                    "filename": filename,
                    "file_size": len(file_content),
                    "lines": 0,
                    "word_count": 0,
                    "file_type": "Text",
                    "parsing_error": str(e),
                },
                "success": False,
                "error": str(e),
            }

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)  # Multiple newlines to double
        text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces/tabs to single space

        # Remove common PDF artifacts
        text = re.sub(r"\x0c", "\n", text)  # Form feed to newline
        text = re.sub(
            r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text
        )  # Control characters

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Strip leading/trailing whitespace
        text = text.strip()

        return text


# Global parser instance
_parser_instance: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """
    Get or create the global document parser instance.

    Returns:
        DocumentParser instance
    """
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Convenience function to extract text from a file.

    Args:
        file_content: Raw file content as bytes
        filename: Name of the file for format detection

    Returns:
        Extracted text content as string
    """
    parser = get_document_parser()
    result = parser.extract_text_from_file(file_content, filename)
    return result["text"]


def extract_text_with_metadata(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract text from a file with detailed metadata.

    Args:
        file_content: Raw file content as bytes
        filename: Name of the file for format detection

    Returns:
        Dictionary containing text, metadata, success status, and error info
    """
    parser = get_document_parser()
    return parser.extract_text_from_file(file_content, filename)
